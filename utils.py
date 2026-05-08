from docx import Document
from io import BytesIO

import streamlit as st


def build_download_text(result_data, edited_cover_letter):
    resume_bullets_text = ""

    for bullet in result_data["resume_bullets"]:
        resume_bullets_text += f"- {bullet}\n"

    suggestions_text = ""

    for suggestion in result_data["improvement_suggestions"]:
        suggestions_text += f"- {suggestion}\n"

    full_text = (
        f"【岗位匹配分析】\n"
        f"{result_data['match_analysis']}\n\n"
        f"【优化后的简历 Bullet Points】\n"
        f"{resume_bullets_text}\n"
        f"【求职信草稿】\n"
        f"{edited_cover_letter}\n\n"
        f"【改进建议】\n"
        f"{suggestions_text}"
    )

    return full_text


def validate_result_data(result_data):
    required_keys = [
        "match_analysis",
        "resume_bullets",
        "cover_letter",
        "improvement_suggestions"
    ]

    for key in required_keys:
        if key not in result_data:
            return False, f"AI 返回结果缺少字段：{key}"

    if not isinstance(result_data["resume_bullets"], list):
        return False, "resume_bullets 应该是一个列表。"

    if not isinstance(result_data["improvement_suggestions"], list):
        return False, "improvement_suggestions 应该是一个列表。"

    return True, ""

def build_file_name(target_job_title,file_type):
    if not target_job_title:
        if file_type=="docx":
            return resume_advice.docx
        return "resume_advice.txt"

    clean_title = target_job_title.strip()
    clean_title = clean_title.replace(" ", "_")
    clean_title = clean_title.replace("/", "_")
    clean_title = clean_title.replace("\\", "_")
    clean_title = clean_title.replace(":", "_")
    if file_type=="docx":
        return f"resume_advice_{clean_title}.docx"
    return f"resume_advice_{clean_title}.txt"


def create_word_file(text):
    document = Document()

    document.add_heading("AI 简历与求职信建议",level=1)
    for line in text.split("\n"):
        document.add_paragraph(line)

    buffer =BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer